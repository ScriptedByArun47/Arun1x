"""
Arun1x Portfolio Backend — v5.0
FastAPI + MongoDB + Cloudinary + Admin Panel
"""

from fastapi import FastAPI, HTTPException, UploadFile, File, Form, Security, Depends, Request
from fastapi.responses import HTMLResponse
from fastapi.security.api_key import APIKeyHeader
from fastapi.middleware.cors import CORSMiddleware
from motor.motor_asyncio import AsyncIOMotorClient
from bson import ObjectId
from datetime import datetime
from typing import Optional
from pydantic import BaseModel
from docx import Document as DocxDocument
from dotenv import load_dotenv
from io import BytesIO
import cloudinary, cloudinary.uploader
import asyncio, os, re

# ─────────────────────────────────────────────
# CONFIG
# ─────────────────────────────────────────────
load_dotenv()

MONGO_URI     = os.getenv("MONGO_URI",     "mongodb://localhost:27017")
DB_NAME       = os.getenv("DB_NAME",       "portfolio_db")
CDN_NAME      = os.getenv("CLOUDINARY_CLOUD_NAME",  "")
CDN_KEY       = os.getenv("CLOUDINARY_API_KEY",      "")
CDN_SECRET    = os.getenv("CLOUDINARY_API_SECRET",   "")
ADMIN_API_KEY = os.getenv("ADMIN_API_KEY", "")

for var, val in [
    ("MONGO_URI", MONGO_URI), ("DB_NAME", DB_NAME),
    ("CLOUDINARY_CLOUD_NAME", CDN_NAME), ("CLOUDINARY_API_KEY", CDN_KEY),
    ("CLOUDINARY_API_SECRET", CDN_SECRET), ("ADMIN_API_KEY", ADMIN_API_KEY),
]:
    if not val:
        raise RuntimeError(f"{var} not set. Check .env / Render env vars.")

cloudinary.config(cloud_name=CDN_NAME, api_key=CDN_KEY, api_secret=CDN_SECRET, secure=True)

print(f"[CONFIG] MongoDB    → {MONGO_URI[:50]}...")
print(f"[CONFIG] DB         → {DB_NAME}")
print(f"[CONFIG] Cloudinary → {CDN_NAME}")

ALLOWED_EXT = {".txt", ".md", ".docx"}
MAX_MB      = 20

# ─────────────────────────────────────────────
# APP  — /docs and /redoc fully disabled
# ─────────────────────────────────────────────
app = FastAPI(
    title       = "Arun1x Portfolio API",
    version     = "5.0.0",
    docs_url    = None,
    redoc_url   = None,
    openapi_url = None,
)

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_methods=["*"],
    allow_headers=["*"],
)

# ─────────────────────────────────────────────
# AUTH
# ─────────────────────────────────────────────
api_key_header = APIKeyHeader(name="X-Admin-Key", auto_error=False)

async def require_admin(key: str = Security(api_key_header)):
    if not key or key != ADMIN_API_KEY:
        raise HTTPException(403, "Forbidden — valid X-Admin-Key header required.")
    return key

# ─────────────────────────────────────────────
# DATABASE
# ─────────────────────────────────────────────
client = db = None

@app.on_event("startup")
async def startup():
    global client, db
    import certifi
    try:
        client = AsyncIOMotorClient(
            MONGO_URI, tls=True,
            tlsCAFile=certifi.where(),
            serverSelectionTimeoutMS=15000,
        )
        db = client[DB_NAME]
        await client.admin.command("ping")
        print(f"[OK] MongoDB connected — {DB_NAME}")
        await db.writeups.create_index("slug", unique=True)
        await db.writeups.create_index("platform")
    except Exception as e:
        print(f"[ERROR] MongoDB: {e}"); raise

@app.on_event("shutdown")
async def shutdown():
    if client: client.close()

def fix_id(doc):
    if doc and "_id" in doc:
        doc["_id"] = str(doc["_id"])
    return doc

def slugify(title):
    s = re.sub(r"[^a-z0-9\s-]", "", title.lower())
    return re.sub(r"\s+", "-", s.strip())[:80]

# ─────────────────────────────────────────────
# CLOUDINARY
# ─────────────────────────────────────────────
def cdn_upload(img_bytes, ext, folder, filename):
    r = cloudinary.uploader.upload(
        BytesIO(img_bytes),
        folder=f"portfolio/writeups/{folder}",
        public_id=filename, resource_type="image",
        format=ext, overwrite=True,
        transformation=[{"quality":"auto","fetch_format":"auto"},{"width":1400,"crop":"limit"}]
    )
    return r["secure_url"]

# ─────────────────────────────────────────────
# PARSERS
# ─────────────────────────────────────────────
def parse_docx(data, slug):
    doc = DocxDocument(BytesIO(data))
    raw = {}
    for rid, rel in doc.part.rels.items():
        if "image" in rel.reltype:
            ext = rel.target_ref.split(".")[-1].lower()
            ext = "jpg" if ext in ("jpg","jpeg") else ext if ext in ("gif","webp","png") else "png"
            raw[rid] = {"bytes": rel.target_part.blob, "ext": ext}

    urls = {}
    for i, (rid, img) in enumerate(raw.items(), 1):
        try:
            urls[rid] = cdn_upload(img["bytes"], img["ext"], slug, f"img_{i}")
            print(f"  [cdn] img_{i} → {urls[rid]}")
        except Exception as e:
            print(f"  [cdn] FAIL img_{i}: {e}"); urls[rid] = ""

    para_img = {}
    for i, para in enumerate(doc.paragraphs):
        for elem in para._element.iter():
            tag = elem.tag.split("}")[-1] if "}" in elem.tag else elem.tag
            if tag == "blip":
                embed = elem.get("{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed")
                if embed and embed in urls:
                    para_img[i] = embed

    lines = []
    for i, para in enumerate(doc.paragraphs):
        if i in para_img:
            u = urls.get(para_img[i], "")
            if u: lines.append(f"![screenshot]({u})")
            lines.append("")
        text  = para.text.strip()
        style = para.style.name if para.style else ""
        if not text: lines.append(""); continue
        if   "Heading 1" in style: lines.append(f"# {text}")
        elif "Heading 2" in style: lines.append(f"## {text}")
        elif "Heading 3" in style: lines.append(f"### {text}")
        elif "Heading 4" in style: lines.append(f"#### {text}")
        elif "List Bullet" in style: lines.append(f"- {text}")
        elif "List Number" in style: lines.append(f"1. {text}")
        elif "Code" in style: lines.append(f"```\n{text}\n```")
        else:
            out = ""
            for run in para.runs:
                t = run.text
                if not t: continue
                if run.bold and run.italic: t = f"***{t}***"
                elif run.bold: t = f"**{t}**"
                elif run.italic: t = f"*{t}*"
                out += t
            lines.append(out.strip() or text)

    for tbl in doc.tables:
        if not tbl.rows: continue
        h = [c.text.strip() for c in tbl.rows[0].cells]
        lines += ["| "+" | ".join(h)+" |", "| "+" | ".join(["---"]*len(h))+" |"]
        for row in tbl.rows[1:]:
            lines.append("| "+" | ".join(c.text.strip() for c in row.cells)+" |")
        lines.append("")

    return "\n".join(lines)

def extract_content(filename, data, slug):
    ext = os.path.splitext(filename)[1].lower()
    if ext == ".docx": return parse_docx(data, slug)
    return data.decode("utf-8", errors="replace")

# ═════════════════════════════════════════════
#  HEALTH
# ═════════════════════════════════════════════
@app.get("/", tags=["Health"])
async def root():
    return {"status": "online", "api": "Arun1x Portfolio API v5.0"}

@app.get("/health-check", tags=["Health"])
async def health_check():
    return {"status": "ok"}

# ═════════════════════════════════════════════
#  PUBLIC API
# ═════════════════════════════════════════════
@app.get("/api/writeups")
async def list_writeups(platform: Optional[str]=None, limit: int=50, skip: int=0):
    q = {}
    if platform and platform != "all": q["platform"] = platform
    cursor = db.writeups.find(q, {"content":0}).sort("created_at",-1).skip(skip).limit(limit)
    docs = [fix_id(d) async for d in cursor]
    return {"writeups": docs, "total": await db.writeups.count_documents(q)}

@app.get("/api/writeups/{slug}")
async def get_writeup(slug: str):
    doc = await db.writeups.find_one({"slug": slug})
    if not doc:
        try: doc = await db.writeups.find_one({"_id": ObjectId(slug)})
        except: pass
    if not doc: raise HTTPException(404, "Writeup not found")
    await db.writeups.update_one({"_id": doc["_id"]}, {"$inc": {"views": 1}})
    return fix_id(doc)

@app.get("/api/projects")
async def get_projects():
    return {"projects": [fix_id(d) async for d in db.projects.find().sort("created_at",-1)]}

@app.get("/api/ctf")
async def get_ctf(platform: Optional[str]=None, status: Optional[str]=None):
    q = {}
    if platform: q["platform"] = platform
    if status:   q["status"]   = status
    return {"challenges": [fix_id(d) async for d in db.ctf.find(q).sort("created_at",-1)]}

@app.get("/api/skills")
async def get_skills():
    return {"skills": [fix_id(d) async for d in db.skills.find().sort("level",-1)]}

@app.get("/api/stats")
async def stats():
    return {
        "writeups":   await db.writeups.count_documents({}),
        "projects":   await db.projects.count_documents({}),
        "ctf_solved": await db.ctf.count_documents({"status":"solved"}),
        "ctf_total":  await db.ctf.count_documents({}),
    }

class ContactIn(BaseModel):
    name: str; email: str; message: str

@app.post("/api/contact", status_code=201)
async def contact(data: ContactIn):
    if len(data.message.strip()) < 10: raise HTTPException(400, "Message too short.")
    if len(data.message) > 2000:       raise HTTPException(400, "Message too long.")
    r = await db.messages.insert_one({**data.dict(), "received_at": datetime.utcnow(), "read": False})
    return {"id": str(r.inserted_id), "status": "received"}

# ═════════════════════════════════════════════
#  ADMIN API  — X-Admin-Key required
# ═════════════════════════════════════════════

# ── Writeup upload ────────────────────────────
@app.post("/api/admin/writeups/upload", status_code=201)
async def upload_writeup(
    file:           UploadFile = File(...),
    title:          str        = Form(...),
    platform:       str        = Form(...),
    platform_label: str        = Form(...),
    excerpt:        str        = Form(...),
    difficulty:     str        = Form(...),
    category:       str        = Form(""),
    tags:           str        = Form(""),
    medium_url:     str        = Form(""),
    _:              str        = Depends(require_admin),
):
    ext = os.path.splitext(file.filename)[1].lower()
    if ext not in ALLOWED_EXT:
        raise HTTPException(400, f"'{ext}' not allowed. Use .txt .md .docx")
    raw = await file.read()
    if len(raw) > MAX_MB * 1024 * 1024:
        raise HTTPException(413, f"File > {MAX_MB}MB")

    base = slugify(title); slug = base; n = 1
    while await db.writeups.find_one({"slug": slug}):
        slug = f"{base}-{n}"; n += 1

    try:
        content = await asyncio.get_event_loop().run_in_executor(
            None, extract_content, file.filename, raw, slug)
    except Exception as e:
        raise HTTPException(422, f"Parse failed: {e}")

    img_count = content.count("![")
    result = await db.writeups.insert_one({
        "title": title, "platform": platform, "platform_label": platform_label,
        "excerpt": excerpt, "difficulty": difficulty, "category": category,
        "tags": [t.strip() for t in tags.split(",") if t.strip()],
        "content": content, "file_name": file.filename, "file_type": ext,
        "slug": slug, "image_count": img_count,
        "medium_url": medium_url,
        "date": datetime.utcnow().strftime("%Y-%m-%d"),
        "created_at": datetime.utcnow(), "views": 0,
    })
    return {"id": str(result.inserted_id), "slug": slug, "image_count": img_count,
            "message": f"Uploaded — {img_count} image(s)"}

@app.delete("/api/admin/writeups/{slug}")
async def delete_writeup(slug: str, _: str = Depends(require_admin)):
    r = await db.writeups.delete_one({"slug": slug})
    if r.deleted_count == 0: raise HTTPException(404, "Not found")
    return {"deleted": True}

# ── Projects ──────────────────────────────────
class ProjectIn(BaseModel):
    title: str; tag: str; description: str
    stack: list; github_url: str = ""; demo_url: str = ""

@app.post("/api/admin/projects", status_code=201)
async def create_project(data: ProjectIn, _: str = Depends(require_admin)):
    r = await db.projects.insert_one({**data.dict(), "created_at": datetime.utcnow()})
    return {"id": str(r.inserted_id)}

@app.delete("/api/admin/projects/{pid}")
async def delete_project(pid: str, _: str = Depends(require_admin)):
    r = await db.projects.delete_one({"_id": ObjectId(pid)})
    if r.deleted_count == 0: raise HTTPException(404, "Not found")
    return {"deleted": True}

# ── CTF ───────────────────────────────────────
class CTFIn(BaseModel):
    name: str; platform: str; category: str
    difficulty: str; status: str; writeup_slug: Optional[str] = None

@app.post("/api/admin/ctf", status_code=201)
async def create_ctf(data: CTFIn, _: str = Depends(require_admin)):
    r = await db.ctf.insert_one({**data.dict(), "created_at": datetime.utcnow()})
    return {"id": str(r.inserted_id)}

@app.delete("/api/admin/ctf/{cid}")
async def delete_ctf(cid: str, _: str = Depends(require_admin)):
    r = await db.ctf.delete_one({"_id": ObjectId(cid)})
    if r.deleted_count == 0: raise HTTPException(404, "Not found")
    return {"deleted": True}

# ── Skills ────────────────────────────────────
class SkillIn(BaseModel):
    name: str; category: str; level: Optional[int] = None
    icon: Optional[str] = None; issuer: Optional[str] = None

@app.post("/api/admin/skills", status_code=201)
async def add_skill(data: SkillIn, _: str = Depends(require_admin)):
    r = await db.skills.insert_one({**data.dict(), "created_at": datetime.utcnow()})
    return {"id": str(r.inserted_id)}

# ── Messages ──────────────────────────────────
@app.get("/api/admin/messages")
async def get_messages(_: str = Depends(require_admin)):
    docs = [fix_id(d) async for d in db.messages.find().sort("received_at",-1)]
    return {"messages": docs, "total": len(docs)}

@app.patch("/api/admin/messages/{mid}/read")
async def mark_read(mid: str, _: str = Depends(require_admin)):
    await db.messages.update_one({"_id": ObjectId(mid)}, {"$set": {"read": True}})
    return {"marked": True}

# ═════════════════════════════════════════════
#  MEDIUM RSS IMPORT
# ═════════════════════════════════════════════
import urllib.request
import xml.etree.ElementTree as ET
import html as html_module

def html_to_markdown(html_str: str, slug: str) -> str:
    """Convert HTML string to Markdown. Uploads images to Cloudinary."""
    from bs4 import BeautifulSoup
    soup = BeautifulSoup(html_str, "html.parser")

    # Use mutable container so nested function can safely append
    state = {"lines": [], "img_counter": 0}

    def process(tag):
        name = getattr(tag, "name", None)
        if name is None:
            t = str(tag).strip()
            if t: state["lines"].append(t)
            return
        if name in ("script","style","nav","footer","button","form"): return
        if name == "h1":
            state["lines"] += [f"# {tag.get_text(strip=True)}", ""]
        elif name == "h2":
            state["lines"] += [f"## {tag.get_text(strip=True)}", ""]
        elif name == "h3":
            state["lines"] += [f"### {tag.get_text(strip=True)}", ""]
        elif name == "h4":
            state["lines"] += [f"#### {tag.get_text(strip=True)}", ""]
        elif name == "p":
            t = tag.get_text(strip=True)
            if t: state["lines"] += [t, ""]
        elif name in ("pre", "code"):
            state["lines"] += [f"```\n{tag.get_text()}\n```", ""]
        elif name == "blockquote":
            state["lines"] += [f"> {tag.get_text(strip=True)}", ""]
        elif name == "hr":
            state["lines"] += ["---", ""]
        elif name in ("ul", "ol"):
            for li in tag.find_all("li", recursive=False):
                state["lines"].append(f"- {li.get_text(strip=True)}")
            state["lines"].append("")
        elif name == "img":
            src = tag.get("src", "") or tag.get("data-src", "")
            alt = tag.get("alt", "screenshot")
            if src and src.startswith("http"):
                state["img_counter"] += 1
                try:
                    req = urllib.request.Request(src, headers={"User-Agent": "Mozilla/5.0"})
                    img_bytes = urllib.request.urlopen(req, timeout=10).read()
                    ext = src.split(".")[-1].split("?")[0][:4] or "jpg"
                    cdn_url = cdn_upload(img_bytes, ext, slug, f"img_{state['img_counter']}")
                    state["lines"] += [f"![{alt}]({cdn_url})", ""]
                    return
                except Exception as e:
                    print(f"  [img] skip: {e}")
                state["lines"] += [f"![{alt}]({src})", ""]
        else:
            for child in tag.children:
                process(child)

    root = soup.body if soup.body else soup
    for child in root.children:
        process(child)

    return "\n".join(state["lines"]).strip()

class MediumRSSIn(BaseModel):
    medium_url:     str            # article URL OR @username profile URL
    platform:       str = "daily"
    platform_label: str = "Medium"
    difficulty:     str = "easy"
    category:       str = ""
    tags:           str = ""

@app.post("/api/admin/import/medium", status_code=201)
async def import_medium_rss(data: MediumRSSIn, _: str = Depends(require_admin)):
    """
    🔒 ADMIN ONLY — Import a Medium article via RSS (no scraping, always works).

    Pass the full article URL:
      https://medium.com/@arun1x/article-slug-xxxx

    The backend extracts the username, fetches the RSS feed,
    finds the matching article, converts HTML → Markdown,
    uploads images to Cloudinary, saves to MongoDB.
    """
    url = data.medium_url.strip()

    # ── Extract Medium username from URL ──
    # Handles: medium.com/@user/slug  OR  user.medium.com/slug
    username = None
    if "medium.com/@" in url:
        username = url.split("medium.com/@")[1].split("/")[0]
    elif ".medium.com" in url:
        username = url.split(".medium.com")[0].split("/")[-1].split(".")[-1]

    if not username:
        raise HTTPException(400, "Could not extract Medium username from URL. "
                               "Use format: https://medium.com/@username/article-slug")

    # ── Fetch RSS feed ──
    rss_url = f"https://medium.com/feed/@{username}"
    print(f"[import] Fetching RSS: {rss_url}")
    try:
        req = urllib.request.Request(
            rss_url,
            headers={"User-Agent": "Mozilla/5.0 (compatible; RSS reader)"}
        )
        rss_bytes = await asyncio.get_event_loop().run_in_executor(
            None, lambda: urllib.request.urlopen(req, timeout=20).read()
        )
    except Exception as e:
        raise HTTPException(400, f"Could not fetch RSS feed for @{username}: {e}")

    # ── Parse XML ──
    try:
        root = ET.fromstring(rss_bytes)
    except ET.ParseError as e:
        raise HTTPException(422, f"RSS parse error: {e}")

    ns = {"content": "http://purl.org/rss/1.0/modules/content/"}
    channel = root.find("channel")
    if not channel:
        raise HTTPException(422, "Invalid RSS feed structure.")

    items = channel.findall("item")
    if not items:
        raise HTTPException(404, f"No articles found in @{username}'s RSS feed.")

    # ── Find matching article by URL ──
    target_item = None
    url_slug = url.rstrip("/").split("/")[-1].lower()

    for item in items:
        item_link = (item.findtext("link") or "").strip()
        item_guid = (item.findtext("guid") or "").strip()
        # Match by slug in URL
        if url_slug and (url_slug in item_link or url_slug in item_guid):
            target_item = item
            break

    # If no match found, use the latest article
    if not target_item:
        target_item = items[0]
        print(f"[import] No URL match — using latest article")

    # ── Extract fields ──
    title   = html_module.unescape(target_item.findtext("title") or "Untitled").strip()
    link    = (target_item.findtext("link") or url).strip()
    pub_date= target_item.findtext("pubDate") or ""
    # Try to parse date
    try:
        from email.utils import parsedate_to_datetime
        date_str = parsedate_to_datetime(pub_date).strftime("%Y-%m-%d")
    except Exception:
        date_str = datetime.utcnow().strftime("%Y-%m-%d")

    # Auto tags from RSS categories
    auto_tags = [c.text for c in target_item.findall("category") if c.text]
    manual_tags = [t.strip() for t in data.tags.split(",") if t.strip()]
    all_tags = list(dict.fromkeys(manual_tags + auto_tags))[:8]

    # Full HTML content from content:encoded
    content_encoded = target_item.find("content:encoded", ns)
    html_content = ""
    if content_encoded is not None and content_encoded.text:
        html_content = content_encoded.text
    else:
        # fallback to description
        html_content = target_item.findtext("description") or ""

    if not html_content:
        raise HTTPException(422, "Article has no content in RSS feed.")

    # Auto excerpt from description
    from bs4 import BeautifulSoup as _BS
    desc_html = target_item.findtext("description") or html_content[:500]
    excerpt = _BS(desc_html, "html.parser").get_text()[:300].strip()

    # ── Build slug ──
    base = slugify(title); sl = base; n = 1
    while await db.writeups.find_one({"slug": sl}):
        sl = f"{base}-{n}"; n += 1

    # ── Convert HTML → Markdown (with Cloudinary image upload) ──
    print(f"[import] Converting HTML → Markdown for '{title}'")
    try:
        content_md = await asyncio.get_event_loop().run_in_executor(
            None, html_to_markdown, html_content, sl
        )
    except Exception as e:
        raise HTTPException(422, f"Content conversion failed: {e}")

    img_count = content_md.count("![")

    # ── Save to MongoDB ──
    result = await db.writeups.insert_one({
        "title":          title,
        "platform":       data.platform,
        "platform_label": data.platform_label,
        "excerpt":        excerpt,
        "difficulty":     data.difficulty,
        "category":       data.category.strip(),
        "tags":           all_tags,
        "content":        content_md,
        "file_name":      link,
        "file_type":      ".rss",
        "slug":           sl,
        "image_count":    img_count,
        "source_url":     link,
        "date":           date_str,
        "created_at":     datetime.utcnow(),
        "views":          0,
    })
    print(f"[import] Saved — slug={sl}, images={img_count}")
    return {
        "id":          str(result.inserted_id),
        "slug":        sl,
        "title":       title,
        "image_count": img_count,
        "tags":        all_tags,
        "message":     f"Imported '{title}' from @{username}'s RSS"
    }

# ═════════════════════════════════════════════
#  ADMIN HTML PANEL  — served at /admin
# ═════════════════════════════════════════════
@app.get("/admin", response_class=HTMLResponse)
async def admin_panel():
    """Serves the admin dashboard HTML page."""
    html_path = os.path.join(os.path.dirname(__file__), "admin.html")
    if not os.path.exists(html_path):
        raise HTTPException(404, "admin.html not found next to main.py")
    with open(html_path, "r") as f:
        return HTMLResponse(content=f.read())

# ═════════════════════════════════════════════
#  DYNAMIC SITEMAP  — for Google indexing
# ═════════════════════════════════════════════
from fastapi.responses import Response
 
@app.get("/sitemap.xml")
async def sitemap():
    base = "https://arun1x.vercel.app"
    urls = [
        f"<url><loc>{base}/</loc><changefreq>weekly</changefreq><priority>1.0</priority></url>",
        f"<url><loc>{base}/index.html</loc><changefreq>weekly</changefreq><priority>0.9</priority></url>",
    ]
    # Add all writeup URLs
    cursor = db.writeups.find({}, {"slug":1, "date":1, "title":1})
    async for doc in cursor:
        slug = doc.get("slug","")
        date = doc.get("date", "")
        urls.append(
            f"<url>"
            f"<loc>{base}/writeup.html?slug={slug}</loc>"
            f"<lastmod>{date}</lastmod>"
            f"<changefreq>monthly</changefreq>"
            f"<priority>0.8</priority>"
            f"</url>"
        )
    xml = '<?xml version="1.0" encoding="UTF-8"?>\n'
    xml += '<urlset xmlns="http://www.sitemaps.org/schemas/sitemap/0.9">\n'
    xml += "\n".join(urls)
    xml += "\n</urlset>"
    return Response(content=xml, media_type="application/xml")

if __name__ == "__main__":
    import uvicorn
    uvicorn.run("main:app", host="0.0.0.0", port=8000, reload=True)
